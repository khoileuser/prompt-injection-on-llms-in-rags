import logging
from pathlib import Path
from typing import Optional

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    Handles loading and processing of documents for prompt injection testing.
    
    This class provides methods to:
    - Load documents from various formats (txt, docx, etc.)
    - Extract text content while preserving special characters
    - Handle document injection attack scenarios
    """
    
    def __init__(self, documents_dir: Optional[str] = None):
        """
        Initialize the document loader.
        
        Args:
            documents_dir: Path to the documents directory. If None,
                          uses the default 'documents' directory relative to
                          the project root.
        """
        if documents_dir is None:
            # Get project root (parent of src directory)
            project_root = Path(__file__).parent.parent
            documents_dir = project_root / "documents"
        
        self.documents_dir = Path(documents_dir)
        
        # Create documents directory if it doesn't exist
        if not self.documents_dir.exists():
            self.documents_dir.mkdir(parents=True, exist_ok=True)
            logger.info(f"Created documents directory: {self.documents_dir}")
        
        logger.info(f"Document loader initialized. Documents directory: {self.documents_dir}")
    
    def load_text_file(self, filename: str) -> Optional[str]:
        """
        Load content from a plain text file.
        
        Args:
            filename: Name of the text file (e.g., 'attack_doc.txt')
            
        Returns:
            String content of the file, or None if file not found
        """
        file_path = self.documents_dir / filename
        
        if not file_path.exists():
            logger.warning(f"Text file not found: {file_path}")
            return None
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            logger.info(f"Loaded text file: {filename} ({len(content)} characters)")
            return content
            
        except Exception as e:
            logger.error(f"Error loading text file {filename}: {str(e)}")
            return None
    
    def load_docx_file(self, filename: str) -> Optional[str]:
        """
        Load content from a Microsoft Word (.docx) file.
        
        Args:
            filename: Name of the docx file (e.g., 'report.docx')
            
        Returns:
            String content of the file, or None if file not found or error
        
        Notes:
            Requires python-docx library. If not installed, falls back gracefully.
        """
        file_path = self.documents_dir / filename
        
        if not file_path.exists():
            logger.warning(f"DOCX file not found: {file_path}")
            return None
        
        try:
            # Try to import python-docx
            try:
                from docx import Document
            except ImportError:
                logger.warning(
                    "python-docx not installed. Install with: pip install python-docx\n"
                    "Falling back to text file loading."
                )
                return None
            
            # Load document
            doc = Document(file_path)
            
            # Extract text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.append(cell.text)
            
            content = '\n'.join(paragraphs)
            
            logger.info(f"Loaded DOCX file: {filename} ({len(content)} characters)")
            return content
            
        except Exception as e:
            logger.error(f"Error loading DOCX file {filename}: {str(e)}")
            return None
    
    def load_document(self, filename: str) -> Optional[str]:
        """
        Load a document of any supported type.
        
        This method automatically detects the file type and uses the
        appropriate loader method.
        
        Args:
            filename: Name of the file to load
            
        Returns:
            String content of the document, or None if loading fails
        """
        if not filename:
            return None
        
        file_ext = Path(filename).suffix.lower()
        
        if file_ext == '.txt':
            return self.load_text_file(filename)
        elif file_ext == '.docx':
            return self.load_docx_file(filename)
        else:
            # Try loading as text file by default
            logger.warning(f"Unknown file extension: {file_ext}, trying as text file")
            return self.load_text_file(filename)
    
    def create_sample_document(self, filename: str, content: str) -> bool:
        """
        Create a sample document for testing purposes.
        
        Args:
            filename: Name of the file to create
            content: Content to write to the file
            
        Returns:
            True if successful, False otherwise
        """
        file_path = self.documents_dir / filename
        
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            logger.info(f"Created sample document: {filename}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating sample document {filename}: {str(e)}")
            return False
    
    def list_documents(self) -> list:
        """
        List all documents in the documents directory.
        
        Returns:
            List of document filenames
        """
        if not self.documents_dir.exists():
            return []
        
        documents = []
        for file_path in self.documents_dir.iterdir():
            if file_path.is_file():
                documents.append(file_path.name)
        
        return sorted(documents)
    
    def get_document_path(self, filename: str) -> Path:
        """
        Get the full path to a document file.
        
        Args:
            filename: Name of the document
            
        Returns:
            Path object pointing to the document
        """
        return self.documents_dir / filename


# =============================================================================
# Global Instance
# =============================================================================

_document_loader: Optional[DocumentLoader] = None


def get_document_loader() -> DocumentLoader:
    """Get the global DocumentLoader instance (singleton pattern)."""
    global _document_loader
    if _document_loader is None:
        _document_loader = DocumentLoader()
    return _document_loader


def load_document(filename: str) -> Optional[str]:
    """
    Convenience function to load a document using the global loader.
    
    Args:
        filename: Name of the document to load
        
    Returns:
        String content of the document, or None if loading fails
    """
    return get_document_loader().load_document(filename)
